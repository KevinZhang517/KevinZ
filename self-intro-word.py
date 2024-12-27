name=input('请输入您的姓名(in Chinese)：')
gender=input('请输入您的性别（男生/女生）：')
age=input('请输入您的年龄：')
motto=input('请输入您的座右铭：')

last_char_name=name[-1]*2
if gender=='男生':
    call='哥哥'
else:call = '宝贝'

if gender=='女生':
    call='姐姐'

if gender!='男生' and gender!="女生":
    call='宝贝'

from docx import Document
doc=Document()
doc.add_heading('自我介绍',0)
doc.add_paragraph(f'我的姓名是{name},你可以叫我{last_char_name}{call}。')
doc.add_paragraph(f'我今年{age}岁')
doc.add_paragraph(f'我的座右铭是:{motto}')


doc.save('个人介绍.docx')

print('\n\n\n\n\n\n')
input("文档已经生成，程序执行完毕，按回车键退出...")

"""

Powered by Kevin Zhang

"""